
# Informações sobre a lib usada para criar word:
# https://python-docx.readthedocs.io/en/latest/user/quickstart.html#applying-a-paragraph-style
## Importante, no caminho do arquivo a barra não pode ser invertida, se não o python não vai entender, como lidar com o caminho

import operator
from operator import concat, contains
from docx import Document
from docx.shared import Inches

file_path = 'C:/Users/luiz.sa/Desktop/ExtrairNotasKindle/My Clippings.txt'
arrayLines, arrayNotes = [], []
dictPositionNotes, lastDict = {}, {}

titleToSearch = input('Insira o nome do título do livro igual está no Kindle: ')
linePosition=0

def getPosition(linePosition):
  print(linePosition)
  getPosition = arrayLines[linePosition+1]
  print(getPosition)
  if "Sua nota" in arrayLines[linePosition+1]:
      position = cleanPositionLine(getPosition)
      return position
  else:
    position = cleanPositionLine(getPosition)
    return position

def getPage(linePage):
  getPageNumber = arrayLines[linePage+1]
  if "Sua nota" in arrayLines[linePage+1]:
      page = getPageNumber[21:26]
  else:
      page = getPageNumber[25:30]
  return page

def cleanPositionLine(getPosition):
  position = ""
  if "Sua nota" in getPosition:
    for characters in getPosition[22:27]:
      if characters.isdigit():
        position = concat(position,characters)
      else:
        break
    return position

  if "destaque" in getPosition:
      for characters in getPosition[26:32]:
        if characters.isdigit():
          position = concat(position,characters)
        else:
          break
      return position

def ordernateNotesNHighlights(position, notes):
    #tempDict = {position:notes} # Comentei para testar a inversão das notas como chaves e posições como valores
    tempDict = {notes:int(position)}
    dictPositionNotes = dict(sorted(tempDict.items(), key=operator.itemgetter(1)))
    return dictPositionNotes

def ifIsNote(lineInformations, lineContent):
  if "Sua nota" in lineInformations:
    newLineContent = concat("== Nota Pessoal abaixo, do próximo destaque: ", lineContent)
    return newLineContent
  else:
    return lineContent

def markRepeatedNotes(dicionario):
  for chaves, valores in dicionario.items():
      for keys, values in dicionario.items():
          if chaves != keys and valores == values:
              if chaves < keys:
                  dicionario[chaves] = "apagar"
              else:
                  dicionario[keys] = "apagar"
  return dicionario

def cleanDictionary(dicionario):
  lastDict = {}
  for chaves, valores in dicionario.items():
      if valores != "apagar":
          lastDict[chaves] = valores
  return lastDict


with open(file_path, encoding='utf8') as clippingsFileObject:
  arrayLines = clippingsFileObject.readlines()

while linePosition < len(arrayLines):
  if titleToSearch in arrayLines[linePosition]:
    position = getPosition(linePosition)
    notes = ifIsNote(arrayLines[linePosition+1], arrayLines[linePosition+3])
    dictPositionNotes.update(ordernateNotesNHighlights(position, notes))          
  linePosition+=1

document = Document()
document.add_heading(titleToSearch)

otherDict = {}
otherDict = dict(sorted(dictPositionNotes.items(), key=operator.itemgetter(1)))
lastDict = markRepeatedNotes(otherDict)
dictFinal = cleanDictionary(lastDict)

for notes in dictFinal.keys():
    if "== Nota Pessoal: " in notes:
      paragraph = document.add_paragraph()
      paragraph.add_run(notes).bold = True
    else:
      document.add_paragraph(notes)

alphanumeric_filter = filter(str.isalnum, titleToSearch)
cleanTitle = "".join(alphanumeric_filter)

print("Separação entre os dicionários\n \n")
print(dictFinal)

document.save('C:/Users/luiz.sa/Desktop/ExtrairNotasKindle/'+cleanTitle+'.docx')

print("O programa foi um sucesso!")

#REFERENCES
#Find and index: https://stackoverflow.com/questions/2294493/how-to-get-the-position-of-a-character-in-python
#Check if String is a digit: https://www.pythonpool.com/python-check-if-string-is-integer/#:~:text=We%20can%20use%20the%20isdigit,Otherwise%2C%20it%20returns%20False.
#Get a sorted copy of a dictionary by value or key: https://stackoverflow.com/questions/613183/how-do-i-sort-a-dictionary-by-value
#Sort Lists: https://www.datacamp.com/community/tutorials/sorting-in-python?utm_source=adwords_ppc&utm_medium=cpc&utm_campaignid=1455363063&utm_adgroupid=65083631748&utm_device=c&utm_keyword=&utm_matchtype=&utm_network=g&utm_adpostion=&utm_creative=278443377095&utm_targetid=aud-392016246653:dsa-429603003980&utm_loc_interest_ms=&utm_loc_physical_ms=9101623&gclid=Cj0KCQiAi9mPBhCJARIsAHchl1xFK3fiPNBMr_p_p9MFTptHeslLvS7pXlYaEZEayXi_ts63BHKeHioaAsFvEALw_wcB
#get value from key: https://note.nkmk.me/en/python-dict-get/#:~:text=You%20can%20use%20get(),the%20key%20does%20not%20exist.&text=Specify%20the%20key%20as%20the,the%20key%20does%20not%20exist.
#Remove white Spaces and non-alphanumeric characters: https://www.kite.com/python/answers/how-to-remove-all-non-alphanumeric-characters-from-a-string-in-python