
# Informações sobre a lib usada para criar word:
# https://python-docx.readthedocs.io/en/latest/user/quickstart.html#applying-a-paragraph-style
## Importante, no caminho do arquivo a barra não pode ser invertida, se não o python não vai entender, como lidar com o caminho

import operator
from operator import concat
from docx import Document
from docx.shared import Inches

file_path = 'C:/Users/luiz.sa/Desktop/ExtrairNotasKindle/My Clippings.txt'
arrayLines, arrayNotes = [], []
dictPositionNotes, tempDict = {}, {}

titleToSearch = input('Insira o nome do título do livro: ')
linePosition=0

def getPosition(linePosition):
  print("Esse foi o linePosition que getPosition() recebeu: ")
  print(linePosition)
  getPosition = arrayLines[linePosition+1]
  print("Esse é o getPosition dentro da função: ")
  print(getPosition)
  if "Sua nota" in arrayLines[linePosition+1]:
      print("Entrou no if, porque tem -sua nota- ")
      position = cleanPositionLine(getPosition)
      print("Chamou cleanPosition e esse é o retorno: ")
      print(position)
      return position
  else:
    position = cleanPositionLine(getPosition)
    print("Chamou cleanPosition e esse é o retorno: ")
    print(position)
    return position

def getPage(linePage):
  getPageNumber = arrayLines[linePage+1]
  if "Sua nota" in arrayLines[linePage+1]:
      page = getPageNumber[21:26]
  else:
      page = getPageNumber[25:30]
  return page

def cleanPositionLine(getPosition):
  print("Aqui é o cleanPosition recebi esse getPosition: ")
  print(getPosition)
  position = ""
  if "Sua nota" in getPosition:
    print("Entrou no if, porque tem -sua nota- ")
    for characters in getPosition[22:27]:
      if characters.isdigit():
        print("Esse caracter é um dígito: " + characters)
        position = concat(position,characters)
        print("Este o position: ")
        print(position)
      else:
        print("Achou o traço e deu Break: ")
        print(position)
        break
    print("Esse é o retorno do cleanPosition de nota: ")
    print(position)
    return position

  if "destaque" in getPosition:
      print("Entrou no if, porque tem -destaque- ")
      for characters in getPosition[26:32]:
        if characters.isdigit():
          print("Esse caracter é um dígito: " + characters)
          position = concat(position,characters)
          print("Este o position: " + position)
        else:
          print("Achou o traço e deu Break: " + position)
          break
      print("Esse é o retorno do cleanPosition de destaque: " + position)
      return position

def ordernateNotesNHighlights(position, notes):
    print("ordenateNotes foi chamado, com esses dados: " + position + " : " + notes)
    #tempDict = {position:notes} # Comentei para testar a inversão das notas como chaves e posições como valores
    tempDict = {notes:int(position)}
    print("Esse é o tempDict dentro da função: ")
    print(tempDict)
    dictPositionNotes = dict(sorted(tempDict.items(), key=operator.itemgetter(1)))
    print("Esse é o dicionário ordenado: ")
    print(dictPositionNotes)
    return dictPositionNotes

def ifIsNote(lineInformations, lineContent):
  print("Se tem - sua nota - nessa linha, adiciona a esse value -Sua nota- : " + lineInformations)
  if "Sua nota" in lineInformations:
    newLineContent = concat("== Nota Pessoal abaixo, do próximo destaque: \n", lineContent)
    print("Se for nota, esse é o novo resultado")
    return newLineContent
  else:
    print("Só retorna o conteúdo" + lineContent)
    return lineContent                      

with open(file_path, encoding='utf8') as clippingsFileObject:
  arrayLines = clippingsFileObject.readlines()

while linePosition < len(arrayLines):
  if titleToSearch in arrayLines[linePosition]:
    print("Chamei a getPosition _____")
    position = getPosition(linePosition)
    print("Esse é o resultado da getPosition para ir pro dicionário: " + position)
    notes = ifIsNote(arrayLines[linePosition+1], arrayLines[linePosition+3])
    print("Essa é a nota final em notes para ir pro dicionário: " + notes)
    dictPositionNotes.update(ordernateNotesNHighlights(position, notes))
    print("Esse é o dicionário final organizado: ")
    print(dictPositionNotes)              
  linePosition+=1

document = Document()
document.add_heading(titleToSearch)

dictFinal = dict(sorted(dictPositionNotes.items(), key=operator.itemgetter(1)))

for notes in dictFinal.keys():
  if "== Nota Pessoal: " in notes:
    paragraph = document.add_paragraph()
    paragraph.add_run(notes).bold = True
  else:
    document.add_paragraph(notes)

alphanumeric_filter = filter(str.isalnum, titleToSearch)
cleanTitle = "".join(alphanumeric_filter)

print(dictPositionNotes)
print("Separação entre os dicionários\n \n")
print(dictFinal)

document.save('C:/Users/luiz.sa/Desktop/ExtrairNotasKindle/'+cleanTitle+'.docx')

#REFERENCES
#Find and index: https://stackoverflow.com/questions/2294493/how-to-get-the-position-of-a-character-in-python
#Check if String is a digit: https://www.pythonpool.com/python-check-if-string-is-integer/#:~:text=We%20can%20use%20the%20isdigit,Otherwise%2C%20it%20returns%20False.
#Get a sorted copy of a dictionary by value or key: https://stackoverflow.com/questions/613183/how-do-i-sort-a-dictionary-by-value
#Sort Lists: https://www.datacamp.com/community/tutorials/sorting-in-python?utm_source=adwords_ppc&utm_medium=cpc&utm_campaignid=1455363063&utm_adgroupid=65083631748&utm_device=c&utm_keyword=&utm_matchtype=&utm_network=g&utm_adpostion=&utm_creative=278443377095&utm_targetid=aud-392016246653:dsa-429603003980&utm_loc_interest_ms=&utm_loc_physical_ms=9101623&gclid=Cj0KCQiAi9mPBhCJARIsAHchl1xFK3fiPNBMr_p_p9MFTptHeslLvS7pXlYaEZEayXi_ts63BHKeHioaAsFvEALw_wcB
#get value from key: https://note.nkmk.me/en/python-dict-get/#:~:text=You%20can%20use%20get(),the%20key%20does%20not%20exist.&text=Specify%20the%20key%20as%20the,the%20key%20does%20not%20exist.
#Remove white Spaces and non-alphanumeric characters: https://www.kite.com/python/answers/how-to-remove-all-non-alphanumeric-characters-from-a-string-in-python
